#!/usr/bin/env python3
"""将 txt 试卷自动排版为符合考试格式的 .docx 文档。"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


def set_run_font(run, font_name: str = "宋体", size: int = 12, bold: bool = False):
    """统一设置 run 字体。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(
    doc: Document,
    text: str,
    font_name: str = "宋体",
    size: int = 12,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    first_line_indent: Cm | None = None,
    space_after: Pt | None = None,
    keep_together: bool = False,
) -> None:
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.keep_together = keep_together
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)


def clean_math(text: str) -> str:
    """将简易 LaTeX 转换为可读文本。"""
    text = text.replace(r"\(", " ").replace(r"\)", " ")
    text = text.replace(r"\left", "").replace(r"\right", "")
    text = re.sub(r"\\frac\{([^}]+)\}\{([^}]+)\}", r"(\1/\2)", text)
    text = text.replace(r"\%", "%")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


HEADER_REGEX = r"(?:一、选择题|二、填空题|三、计算题|四、操作题|五、解决问题)"


def normalize_input(text: str) -> str:
    """预处理：清理 LaTeX、把挤在一行的多个小题拆开。"""
    # 逐行清理数学公式，保留换行结构
    lines = [clean_math(line) for line in text.splitlines()]
    text = "\n".join(lines)

    # 把 ①②③... 前面加换行
    text = re.sub(r"([①②③④⑤⑥⑦⑧⑨⑩])", r"\n\1", text)

    # 把 "1. xxx 2. xxx" 这种子标题前面加换行
    text = re.sub(r"(\D)(\d+[\.．、])(?=\D)", r"\1\n\2", text)

    # 把连在一行的选项 "A. x  B. y  C. z  D. w" 拆成多行
    # 匹配 "A. " 后面跟着内容，再跟两个以上空格和 "B." 等
    text = re.sub(r"([A-D][\.．、])(.*?)(\s{2,})(?=[A-D][\.．、])", r"\1\2\n", text)

    # 分值标记后如果紧跟文字，加换行
    text = re.sub(r"(（\s*\d+\s*分\s*）)(?=[^\n])", r"\1\n", text)

    # 清理空行
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def split_sections(text: str) -> tuple[str, str, list[dict]]:
    """切分卷头信息与各大题。"""
    lines = text.splitlines()
    first_line = lines[0] if lines else ""

    title_match = re.match(r"^(.*?卷)(.*)$", first_line.strip())
    title = first_line.strip()
    extra_info = ""
    if title_match:
        title = title_match.group(1).strip()
        extra_info = title_match.group(2).strip()

    # 如果 extra_info 里还包含大题标题，把大题标题部分还回去
    m = re.search(HEADER_REGEX, extra_info)
    if m:
        extra_info = extra_info[: m.start()].strip()

    full_text = "\n".join(lines)
    positions = [m.start() for m in re.finditer(HEADER_REGEX, full_text)]

    sections = []
    if positions:
        for i, start in enumerate(positions):
            end = positions[i + 1] if i + 1 < len(positions) else len(full_text)
            section_text = full_text[start:end].strip()
            header_match = re.match(r"(" + HEADER_REGEX + r")(.*)", section_text, re.DOTALL)
            if header_match:
                header = header_match.group(1)
                body = header_match.group(2).strip()
                sections.append({"header": header, "body": body})

    return title, extra_info, sections


def strip_score(text: str) -> str:
    """去掉题干前的大题分值说明。"""
    return re.sub(r"^\s*（[^）]*分[^）]*）\s*", "", text).strip()


def split_by_score_markers(text: str) -> list[str]:
    """按（X 分）切分解决问题/操作题的小题。"""
    parts = re.split(r"(（\s*\d+\s*分\s*）)", text)
    items = []
    i = 0
    while i < len(parts):
        part = parts[i].strip()
        if not part:
            i += 1
            continue
        if i + 1 < len(parts) and re.match(r"（\s*\d+\s*分\s*）", parts[i + 1]):
            part = part + parts[i + 1]
            i += 2
        else:
            i += 1
        if part:
            items.append(part)
    return items


def is_option_line(line: str) -> bool:
    return bool(re.match(r"^[A-D][\.．、]", line.strip()))


def is_question_stem(line: str) -> bool:
    """判断一行是否像题干（以全角括号结尾或包含填空括号）。"""
    line = line.strip()
    return line.endswith("（）") or "（）" in line or "？" in line[-3:]


def split_option_line(line: str) -> list[str]:
    """把连在一行的多个选项拆成单独选项。"""
    line = line.strip()
    # 在 B./C./D. 前面加换行，要求前面是空白或选项开头
    line = re.sub(r"(?<=\s)([B-D][\.．、])", r"\n\1", line)
    return [opt.strip() for opt in line.splitlines() if opt.strip()]


def parse_choice_questions(body: str) -> list[dict]:
    """解析选择题。"""
    lines = [line.strip() for line in body.splitlines() if line.strip()]
    questions = []
    current: dict | None = None

    for line in lines:
        if is_option_line(line):
            if current is None:
                current = {"stem": "", "options": []}
            current["options"].extend(split_option_line(line))
            continue

        if current is None:
            current = {"stem": line, "options": []}
            continue

        # 新题干开始：当前题已有选项，或当前题无选项但新行也是题干
        if current.get("options") and is_question_stem(line):
            questions.append(current)
            current = {"stem": line, "options": []}
        elif not current.get("options") and is_question_stem(line):
            questions.append(current)
            current = {"stem": line, "options": []}
        else:
            current["stem"] += line

    if current:
        questions.append(current)

    return questions


def parse_questions(section: dict) -> list[dict]:
    """解析单个大题内的小题。"""
    body = strip_score(section["body"])
    header = section["header"]
    questions = []

    if "选择" in header:
        questions = parse_choice_questions(body)

    elif "填空" in header:
        for line in body.splitlines():
            line = line.strip()
            if line:
                questions.append({"stem": line})

    elif "计算" in header:
        # 按 1. / 2. 切分子类型，再按 ①②③ 切分小题
        sub_sections = re.split(r"\n(?=\d+[\.．、])", body)
        for sub in sub_sections:
            sub = sub.strip()
            if not sub:
                continue
            m = re.match(r"(\d+[\.．、])([^①②③④⑤⑥⑦⑧⑨⑩]*)(.*)", sub, re.DOTALL)
            if m:
                sub_title = m.group(2).strip()
                items_text = m.group(3).strip()
                items = [it.strip() for it in re.split(r"(?=[①②③④⑤⑥⑦⑧⑨⑩])", items_text) if it.strip()]
                questions.append({"stem": sub_title, "items": items})
            else:
                items = [it.strip() for it in re.split(r"(?=[①②③④⑤⑥⑦⑧⑨⑩])", sub) if it.strip()]
                if items:
                    questions.append({"stem": "", "items": items})

    elif "操作" in header or "解决问题" in header:
        for item in split_by_score_markers(body):
            item = strip_score(item).strip()
            if item:
                questions.append({"stem": item})

    else:
        for line in body.splitlines():
            line = line.strip()
            if line:
                questions.append({"stem": line})

    return questions


def build_docx(input_path: Path, output_path: Path) -> Path:
    doc = Document()

    # 页面设置：A4，标准页边距
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    raw_text = input_path.read_text(encoding="utf-8")
    normalized_text = normalize_input(raw_text)
    title, extra_info, sections = split_sections(normalized_text)

    # 标题
    add_paragraph(doc, title, font_name="黑体", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    if extra_info:
        add_paragraph(doc, extra_info, font_name="宋体", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # 考生信息栏
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "姓名"
    hdr_cells[1].text = "班级"
    hdr_cells[2].text = "考号"
    hdr_cells[3].text = "得分"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, "宋体", 12, bold=True)
    doc.add_paragraph()

    # 各大题
    for sec in sections:
        add_paragraph(doc, sec["header"], font_name="黑体", size=14, bold=True, space_after=Pt(6))

        questions = parse_questions(sec)
        for idx, q in enumerate(questions, 1):
            if "选择" in sec["header"]:
                add_paragraph(doc, f"{idx}. {q['stem']}", first_line_indent=Cm(0), space_after=Pt(3))
                if q.get("options"):
                    # 每个选项单独一行，方便排版
                    options_text = "    ".join(q["options"])
                    add_paragraph(doc, options_text, first_line_indent=Cm(0.74), space_after=Pt(6))
            elif "计算" in sec["header"] and q.get("items"):
                add_paragraph(doc, f"{idx}. {q['stem']}", space_after=Pt(3))
                for item in q["items"]:
                    add_paragraph(doc, item, first_line_indent=Cm(0.74), space_after=Pt(3))
            else:
                add_paragraph(doc, f"{idx}. {q['stem']}", first_line_indent=Cm(0), space_after=Pt(6))

    doc.save(output_path)
    return output_path


def main() -> int:
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("docs/试卷.txt")
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("data/write-output/试卷-排版.docx")

    if not input_path.exists():
        print(f"输入文件不存在：{input_path}", file=sys.stderr)
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    build_docx(input_path, output_path)
    print(f"已生成考试格式文档：{output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
